"""
Word Document Editing Service for Schengen Visa Application Forms
"""

import os
import tempfile
from typing import Dict, Any, Optional
from docx import Document
from docx.shared import Inches
import logging

logger = logging.getLogger(__name__)


class WordDocumentService:
    """Service for editing Word documents with form data"""
    
    def __init__(self):
        # Get the directory where this service file is located
        current_dir = os.path.dirname(os.path.abspath(__file__))
        self.template_path = os.path.join(current_dir, "schengen-visa-application-form_ocred.docx")
        
        # Set output directory to backend/generated_documents
        backend_dir = os.path.dirname(os.path.dirname(current_dir))
        self.output_dir = os.path.join(backend_dir, "generated_documents")
        
        # Ensure output directory exists
        os.makedirs(self.output_dir, exist_ok=True)
        
        logger.info(f"Template path: {self.template_path}")
        logger.info(f"Output directory: {self.output_dir}")
    
    def edit_document(self, user_data: Dict[str, Any], filename: Optional[str] = None) -> str:
        """
        Edit the Word document with user data
        
        Args:
            user_data: Dictionary containing form field data
            filename: Optional custom filename for the output
            
        Returns:
            Path to the generated document
        """
        try:
            # Check if template exists
            if not os.path.exists(self.template_path):
                raise FileNotFoundError(f"Template file not found: {self.template_path}")
            
            # Load the document
            doc = Document(self.template_path)
            
            logger.info(f"Document has {len(doc.paragraphs)} paragraphs and {len(doc.tables)} tables")
            logger.info(f"Received {len(user_data)} fields to replace")
            logger.info(f"Fields: {list(user_data.keys())}")
            
            # Process paragraphs
            self._process_paragraphs(doc, user_data)
            
            # Process tables
            self._process_tables(doc, user_data)
            
            # Generate output filename
            if not filename:
                filename = f"schengen-visa-application-form_filled_{self._generate_timestamp()}.docx"
            
            output_path = os.path.join(self.output_dir, filename)
            
            # Save the modified document
            doc.save(output_path)
            logger.info(f"Document saved to: {output_path}")
            
            return output_path
            
        except Exception as e:
            logger.error(f"Error editing Word document: {e}")
            raise
    
    def _process_paragraphs(self, doc: Document, user_data: Dict[str, Any]) -> None:
        """Process paragraphs in the document"""
        replacements_made = 0
        for paragraph in doc.paragraphs:
            # Process each run to preserve formatting
            for run in paragraph.runs:
                for field_key, field_value in user_data.items():
                    if field_key in run.text:
                        run.text = run.text.replace(field_key, str(field_value))
                        replacements_made += 1
                        logger.info(f"Replaced {field_key} with '{field_value}' in paragraph")
        
        logger.info(f"Made {replacements_made} replacements in paragraphs")
    
    def _process_tables(self, doc: Document, user_data: Dict[str, Any]) -> None:
        """Process tables in the document"""
        replacements_made = 0
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    # Process each paragraph in the cell
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            for field_key, field_value in user_data.items():
                                if field_key in run.text:
                                    run.text = run.text.replace(field_key, str(field_value))
                                    replacements_made += 1
                                    logger.info(f"Replaced {field_key} with '{field_value}' in table")
        
        logger.info(f"Made {replacements_made} replacements in tables")
    
    def _generate_timestamp(self) -> str:
        """Generate timestamp for unique filenames"""
        from datetime import datetime
        return datetime.now().strftime("%Y%m%d_%H%M%S")
    
    def get_sample_data(self) -> Dict[str, Any]:
        """Get sample data for testing"""
        return {
            "FIELD1": "DOE",  # SURNAME
            "FIELD2": "SMITH",  # Surname at birth
            "FIELD3": "John",  # First Names
            "FIELD4": "15/03/1990",  # Date of Birth
            "FIELD5": "New York",  # Place of birth
            "FIELD6": "United States",  # Country of birth
            "FIELD7": "American",  # Current Nationality
            "FIELD8": "Jane Doe, 123 Main St, New York, +1-555-123-4567, jane@example.com, American",  # Parental Authority
            "FIELD9": "123456789",  # National identity number
            "FIELD10": "DOE",  # Surname (family name) of family member
            "FIELD11": "Jane",  # First name of family member
            "FIELD12": "20/05/1995",  # Birth of family member
            "FIELD13": "American",  # Nationality of family member
            "FIELD14": "987654321",  # Number of ID of family member
            "FIELD15": "123 Main Street, New York, NY 10001, john.doe@example.com",  # Address and email
            "FIELD16": "+1-555-123-4567",  # Telephone number
            "FIELD17": "Software Engineer",  # Current occupation
            "FIELD18": "Tech Corp, 456 Business Ave, New York, NY 10002, +1-555-987-6543",  # Employer's address and phone
            "FIELD19": "Tourism",  # Purpose of stay
            "FIELD20": "Germany",  # Destination Country
            "FIELD21": "Hotel Berlin, 123 Tourist Street, Berlin, Germany",  # Address of accommodation
            "FIELD22": "+49-30-12345678",  # Tel no of accommodation
            "FIELD23": "Business Solutions GmbH",  # Inviting company's name
            "FIELD24": "Germany",  # Member State of main destination
            "FIELD25": "Germany",  # Member State of first entry
            "FIELD26": "US123456789",  # Number of travel document
            "FIELD27": "01/01/2020",  # Date of issue
            "FIELD28": "01/01/2030",  # Valid until
            "FIELD29": "United States",  # Issued by Country
            "FIELD30": "Hans Mueller",  # Surname and first name of inviting person
            "FIELD31": "Maria Schmidt, 456 Business St, Berlin, +49-30-98765432, maria@business.de",  # Contact person details
            "FIELD32": "+49-30-11111111",  # Telephone no. of company/organisation
            "FIELD33": "Self-funded",  # COST
            "FIELD34": "Germany, 15/06/2025",  # Country and Date
        }
    
    def validate_user_data(self, user_data: Dict[str, Any]) -> Dict[str, Any]:
        """Validate and clean user data"""
        validated_data = {}
        
        for field_num in range(1, 35):
            field_key = f"FIELD{field_num}"
            value = user_data.get(field_key, "")
            
            # Basic validation and cleaning
            if isinstance(value, str):
                validated_data[field_key] = value.strip()
            else:
                validated_data[field_key] = str(value) if value is not None else ""
        
        return validated_data
